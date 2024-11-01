import streamlit as st
import numpy as np
import sounddevice as sd
import soundfile as sf
import speech_recognition as sr
from docx import Document
import os

# Tạo thư mục để lưu file
UPLOAD_FOLDER = 'uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# Hàm chuyển đổi âm thanh thành văn bản
def convert_audio_to_text(audio_path):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio = recognizer.record(source)
    try:
        text = recognizer.recognize_google(audio, language='vi-VN')
    except sr.UnknownValueError:
        text = "Không thể nhận diện được giọng nói"
    except sr.RequestError:
        text = "Lỗi khi yêu cầu dịch vụ nhận diện giọng nói"
    return text

# Hàm tạo tài liệu Word
def create_word_document(text, output_path):
    doc = Document()
    doc.add_heading('Transcript', level=1)
    doc.add_paragraph(text)
    doc.save(output_path)

# Ghi âm từ microphone
def record_audio(duration=5):
    st.write("Ghi âm... Nhấn nút dừng để kết thúc.")
    audio_data = sd.rec(int(duration * 44100), samplerate=44100, channels=1, dtype='float64')
    sd.wait()  # Chờ cho đến khi ghi âm hoàn tất
    audio_file = 'output.wav'
    sf.write(audio_file, audio_data, 44100)  # Lưu file âm thanh
    return audio_file

# Giao diện người dùng Streamlit
st.title("Chuyển Đổi Âm Thanh Thành Văn Bản")
st.write("Ghi âm từ microphone hoặc tải lên file âm thanh.")

# Chọn ghi âm từ microphone
if st.button("Ghi âm từ Microphone"):
    audio_file_path = record_audio()
    st.success("Ghi âm hoàn tất!")

    if st.button("Chuyển đổi thành văn bản"):
        text = convert_audio_to_text(audio_file_path)
        st.write("Kết quả:", text)

        # Tạo tài liệu Word
        output_path = os.path.join(UPLOAD_FOLDER, 'transcript.docx')
        create_word_document(text, output_path)
        st.download_button("Tải tài liệu Word", output_path, file_name='transcript.docx')

# Tải lên file âm thanh
uploaded_file = st.file_uploader("Tải lên file âm thanh", type=["wav", "mp3", "ogg"])
if uploaded_file is not None:
    audio_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(audio_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    st.success("Tải lên thành công!")

    if st.button("Chuyển đổi thành văn bản"):
        text = convert_audio_to_text(audio_path)
        st.write("Kết quả:", text)

        # Tạo tài liệu Word
        output_path = os.path.join(UPLOAD_FOLDER, 'transcript.docx')
        create_word_document(text, output_path)
        st.download_button("Tải tài liệu Word", output_path, file_name='transcript.docx')
